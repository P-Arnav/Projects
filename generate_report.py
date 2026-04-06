"""
Generates the project report Word document.
Run once: python generate_report.py
Output:   MV_Project_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code(doc, text):
    """Add a monospaced code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()   # spacing after table
    return table


def add_spacer(doc):
    doc.add_paragraph()


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "A History-Conditioned Cross-Modal Attention Framework\n"
    "for Multimodal Lung Disease Diagnosis from Chest X-Rays"
)
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Project Report — VIT University").font.size = Pt(13)

doc.add_paragraph()
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run("Based on: Arnav Panwar & David Kumar").font.size = Pt(11)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", level=1)

add_body(doc,
    "Chest X-ray interpretation is one of the most frequently performed yet error-prone "
    "tasks in radiology. Automated deep learning systems have made significant progress "
    "in image-based diagnosis, but they share a fundamental limitation with naive human "
    "reading: they look only at the image and ignore all other available clinical context."
)
add_spacer(doc)
add_body(doc,
    "In real clinical practice, a radiologist never reads an image in isolation. Patient "
    "age, sex, prior visit history, and the projection view of the X-ray all influence the "
    "interpretation. A shadow that might suggest cardiomegaly in a 70-year-old male looks "
    "very different from the same finding in a 25-year-old athlete. This project implements "
    "a multimodal architecture that incorporates structured patient history alongside the "
    "chest X-ray image, fusing the two modalities through bidirectional cross-modal attention."
)
add_spacer(doc)

add_heading(doc, "1.1 Objectives", level=2)
add_bullet(doc, "Implement a full multimodal pipeline for multi-label chest X-ray diagnosis.")
add_bullet(doc, "Fuse visual features (DenseNet-121) with clinical history text (Bio_ClinicalBERT) using bidirectional cross-modal attention.")
add_bullet(doc, "Introduce a consistency module that gates predictions based on how well the image and history agree.")
add_bullet(doc, "Demonstrate the complete system on the NIH ChestXray14 benchmark (112,120 images, 14 disease classes).")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. DATASET
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Dataset: NIH ChestXray14", level=1)

add_body(doc,
    "The NIH ChestXray14 dataset is the standard public benchmark for automated chest X-ray "
    "classification. Released by Wang et al. (2017), it contains frontal-view X-ray images "
    "with weak labels mined from radiology reports using NLP."
)
add_spacer(doc)

add_heading(doc, "2.1 Dataset Statistics", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        ["Total images",          "112,120"],
        ["Unique patients",       "30,805"],
        ["Disease classes",       "14 (multi-label)"],
        ["Label format",          "Binary vector per image (one entry per class)"],
        ["Image format",          "PNG, 1024×1024 pixels (resized to 224×224 for training)"],
        ["Train / Val split",     "80% / 20% of official train_val_list.txt"],
        ["Official test set",     "test_list.txt (held out)"],
    ]
)

add_heading(doc, "2.2 Disease Classes", level=2)
add_body(doc,
    "The 14 pathology classes are: Atelectasis, Cardiomegaly, Effusion, Infiltration, "
    "Mass, Nodule, Pneumonia, Pneumothorax, Consolidation, Edema, Emphysema, Fibrosis, "
    "Pleural Thickening, and Hernia. Each image may carry zero or more labels simultaneously "
    "(multi-label classification). A special 'No Finding' label is used when none of the 14 "
    "pathologies are present, but it is excluded from the prediction target vector."
)
add_spacer(doc)

add_heading(doc, "2.3 Synthetic Clinical History", level=2)
add_body(doc,
    "The NIH dataset provides images and disease labels but contains no free-text clinical "
    "notes. To enable multimodal training, we generate a synthetic history string from the "
    "available metadata fields for each image:"
)
add_bullet(doc, "Patient age")
add_bullet(doc, "Patient gender (M / F)")
add_bullet(doc, "View position (PA = posterior-anterior, AP = anterior-posterior)")
add_bullet(doc, "Follow-up number (0 = initial presentation, >0 = return visit)")
add_bullet(doc, "Disease labels from the patient's previous visit (if follow-up > 0)")
add_spacer(doc)
add_body(doc, "Example generated history:")
add_code(doc,
    "\"Patient is a 67-year-old male. Chest X-ray taken in PA view. "
    "Follow-up visit number 3. Previous findings: Effusion, Atelectasis.\""
)
add_spacer(doc)
add_body(doc,
    "When required metadata is absent (e.g. age = 0), the special token 'NO_HISTORY' is "
    "used. The text encoder is trained to produce a neutral embedding for this sentinel value."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3. ARCHITECTURE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Model Architecture", level=1)

add_body(doc,
    "The full model is called MultimodalDiagnosticModel and consists of four sequential "
    "stages: visual encoding, text encoding, cross-modal fusion, and consistency-gated "
    "classification. The overall data flow is:"
)
add_spacer(doc)
add_code(doc,
    "X-Ray Image  →  DenseNet-121  →  Visual Tokens  (B, 49, 1024)\n"
    "History Text →  Bio_ClinicalBERT →  Text Tokens (B, N_t, 768)\n"
    "                    ↓\n"
    "         Bidirectional Cross-Modal Attention  (shared 512-dim)\n"
    "                    ↓\n"
    "         Consistency Gate (MLP compatibility score s)\n"
    "                    ↓\n"
    "         14-class sigmoid output → disease probabilities"
)
add_spacer(doc)

# 3.1 Visual Encoder
add_heading(doc, "3.1 Visual Encoder — DenseNet-121", level=2)
add_body(doc,
    "DenseNet-121 (Huang et al., 2017) is a 121-layer convolutional neural network with "
    "dense connectivity: every layer receives feature maps from all preceding layers within "
    "its dense block. This design has two key advantages for medical imaging:"
)
add_bullet(doc, "Feature reuse: earlier low-level edge and texture features are directly accessible to deeper layers, which is important for detecting subtle pathologies.")
add_bullet(doc, "Gradient flow: skip connections allow gradients to propagate directly to early layers, preventing vanishing gradients in deep networks.")
add_spacer(doc)
add_body(doc, "Architecture details:")
add_table(doc,
    ["Component", "Detail"],
    [
        ["Input",             "224×224×3 RGB image (ImageNet normalisation)"],
        ["Dense blocks",      "4 dense blocks with 6, 12, 24, 16 layers respectively"],
        ["Transition layers", "1×1 conv + 2×2 avg pool between dense blocks"],
        ["Final feature map", "(B, 1024, 7, 7) — 1024 channels, 7×7 spatial grid"],
        ["Token sequence",    "Flatten spatial grid → (B, 49, 1024) region tokens"],
        ["Projection",        "Linear(1024 → 512) inside the fusion module"],
        ["Pretraining",       "ImageNet-1K weights (torchvision DenseNet121_Weights.IMAGENET1K_V1)"],
        ["Fine-tuning",       "All layers trainable with backbone_lr = 1e-5"],
    ]
)
add_body(doc,
    "Note: DenseNet-121 was chosen because CheXNet (Rajpurkar et al., 2017) first demonstrated "
    "its superiority over radiologists on the NIH CXR14 benchmark, making it the established "
    "backbone for this task. The final 7×7 spatial feature map is kept as a token sequence "
    "(49 region tokens) rather than global average pooled, so the cross-modal attention can "
    "attend to specific spatial regions of the lung."
)
add_spacer(doc)

# 3.2 Text Encoder
add_heading(doc, "3.2 Text Encoder — Bio_ClinicalBERT", level=2)
add_body(doc,
    "Bio_ClinicalBERT (Alsentzer et al., 2019) is a domain-adapted BERT model pretrained on "
    "clinical notes from the MIMIC-III database. It shares the BERT-base architecture but has "
    "been further pretrained on approximately 880 million words of de-identified hospital "
    "electronic health records, giving it strong representations of medical terminology."
)
add_spacer(doc)
add_table(doc,
    ["Component", "Detail"],
    [
        ["Base architecture",   "BERT-base (bert-base-uncased architecture)"],
        ["Transformer layers",  "12 encoder layers"],
        ["Hidden dimension",    "768"],
        ["Attention heads",     "12 per layer"],
        ["Parameters",          "~110 million"],
        ["Pretraining corpus",  "MIMIC-III clinical notes (~880M words)"],
        ["HuggingFace ID",      "emilyalsentzer/Bio_ClinicalBERT"],
        ["Input format",        "[CLS] <history text> [SEP], padded to max_length=128"],
        ["Output used",         "[CLS] token embedding (768-dim) as sentence representation"],
        ["Token embeddings",    "Full sequence (B, N_t, 768) used for cross-modal attention"],
        ["Projection",          "Linear(768 → 512) inside the fusion module"],
        ["Fine-tuning",         "Layers 0–7 frozen; layers 8–11 trainable with backbone_lr = 1e-5"],
    ]
)
add_body(doc,
    "The [CLS] token is a special prepended token whose final hidden state serves as an "
    "aggregated sentence-level representation. It is used by the consistency module. The "
    "full per-token sequence is used by the cross-modal attention to allow fine-grained "
    "interaction between specific history tokens (e.g. 'male', '67', 'PA') and specific "
    "spatial regions of the X-ray."
)
add_spacer(doc)

# 3.3 Lung ROI
add_heading(doc, "3.3 Lung ROI Segmentation — torchxrayvision PSPNet", level=2)
add_body(doc,
    "Before the image is passed to DenseNet-121, an optional lung region-of-interest (ROI) "
    "crop is applied. The PSPNet segmentation model from the torchxrayvision library segments "
    "the lung fields and crops the image to the bounding box of the detected lung region. "
    "This removes irrelevant background (shoulders, table edges) and focuses the visual "
    "encoder on the clinically relevant area. The cropped region is then resized to 224×224 "
    "for input to DenseNet-121."
)
add_spacer(doc)

# 3.4 Cross-Modal Fusion
add_heading(doc, "3.4 Cross-Modal Attention Fusion", level=2)
add_body(doc,
    "The fusion module implements bidirectional token-level cross-modal attention in a shared "
    "512-dimensional embedding space. Both modality token sequences are first projected into "
    "this shared space via separate linear layers:"
)
add_spacer(doc)
add_code(doc, "V' = Linear(1024 → 512)(V)    # visual tokens:  (B, 49, 512)")
add_code(doc, "T' = Linear(768  → 512)(T)    # text tokens:    (B, N_t, 512)")
add_spacer(doc)
add_body(doc, "Two cross-attention operations are then computed in parallel:")
add_spacer(doc)
add_bullet(doc, "Visual attends to text (V→T): each visual token queries the text sequence. This lets each spatial region of the X-ray look up relevant history tokens (e.g., the cardiac region can attend to the word 'male' or '70').", bold_prefix="V→T attention:  ")
add_bullet(doc, "Text attends to visual (T→V): each history token queries the visual token sequence. This lets the word 'emphysema' in the prior history focus on the upper lung region where emphysema typically appears.", bold_prefix="T→V attention:  ")
add_spacer(doc)
add_body(doc,
    "Each cross-attention uses 8 heads, with head dimension 64 (512 / 8). The output of "
    "each direction is added back to the input via a residual connection and layer-normalised, "
    "following the standard transformer pattern. The T→V attention weights are also returned "
    "as output for explainability (attention heatmaps)."
)
add_spacer(doc)

# 3.5 Consistency Module
add_heading(doc, "3.5 Consistency Module", level=2)
add_body(doc,
    "After fusion, a consistency module computes a scalar compatibility score s ∈ [0, 1] "
    "that measures how well the visual and textual representations agree. The score is used "
    "to gate the contribution of the fused representation in the final prediction."
)
add_spacer(doc)
add_body(doc, "The module takes three inputs:")
add_bullet(doc, "V_fused — mean-pooled fused visual tokens (B, 512)")
add_bullet(doc, "T_fused — mean-pooled fused text tokens (B, 512)")
add_bullet(doc, "t_cls   — [CLS] embedding from ClinicalBERT (B, 768), as a global history summary")
add_spacer(doc)
add_body(doc, "These are concatenated and passed through an MLP:")
add_code(doc, "h = concat([V_fused, T_fused, t_cls])   # (B, 512 + 512 + 768) = (B, 1792)")
add_code(doc, "s = sigmoid(MLP(h))                     # (B, 1),  s ∈ [0, 1]")
add_spacer(doc)
add_body(doc,
    "The final prediction combines an image-only branch and a fused branch, gated by s:"
)
add_code(doc, "logits_img   = head_img(proj(V.mean(dim=1)))          # image-only path")
add_code(doc, "logits_fused = head_fused((V_fused * s).mean(dim=1))  # fused path, scaled by s")
add_code(doc, "logits       = logits_img + s * logits_fused           # gated combination (eq. 15)")
add_spacer(doc)
add_body(doc,
    "When the history is missing or uninformative (s → 0), the model falls back to the "
    "image-only prediction. When image and history strongly agree (s → 1), the fused "
    "prediction is fully incorporated."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4. TRAINING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Training", level=1)

add_heading(doc, "4.1 Loss Function", level=2)
add_body(doc,
    "The training objective combines a primary classification loss with a consistency "
    "regularisation term:"
)
add_spacer(doc)
add_code(doc, "L = L_cls + λ · L_cons")
add_spacer(doc)
add_bullet(doc, "Binary Cross-Entropy with Logits applied independently to each of the 14 disease classes. A per-class positive weight (neg_count / pos_count, clamped to 50) is used to handle severe label imbalance.", bold_prefix="L_cls:  ")
add_bullet(doc, "Encourages agreement between the gated output and the image-only output, weighted by the consistency score. This prevents the model from ignoring one modality entirely.", bold_prefix="L_cons: ")
add_bullet(doc, "Consistency loss weight, set to 0.1 as per the paper.", bold_prefix="λ:      ")
add_spacer(doc)

add_heading(doc, "4.2 Optimiser and Schedule", level=2)
add_table(doc,
    ["Hyperparameter", "Value", "Notes"],
    [
        ["Optimiser",          "AdamW",      "Weight decay = 1e-4"],
        ["New layer LR",       "1e-4",       "Cross-attention, consistency module, heads"],
        ["Backbone LR",        "1e-5",       "DenseNet-121 and ClinicalBERT fine-tuning"],
        ["LR schedule",        "Cosine annealing", "T_max = num_epochs, eta_min = 1e-6"],
        ["Gradient clipping",  "max_norm = 1.0",   "Applied before each optimizer.step()"],
        ["Precision",          "FP32 (default)",   "FP16 mixed precision optional"],
        ["Epochs",             "50",               "Default; early stopping via best AUC"],
        ["Batch size",         "32",               "Adjust based on GPU memory"],
    ]
)

add_heading(doc, "4.3 Class Imbalance", level=2)
add_body(doc,
    "The NIH dataset is highly imbalanced — Hernia appears in fewer than 0.2% of images "
    "while Infiltration appears in over 17%. Uniform BCE would cause the model to predict "
    "all negatives for rare classes. We address this with a per-class positive weight "
    "computed over the training set before training begins:"
)
add_spacer(doc)
add_code(doc, "pos_weight[c] = neg_count[c] / pos_count[c],   clipped to [0, 50]")
add_spacer(doc)
add_body(doc,
    "This weight is passed directly to PyTorch's BCEWithLogitsLoss, which scales the "
    "gradient contribution of positive samples by this factor."
)
add_spacer(doc)

add_heading(doc, "4.4 Differential Learning Rates", level=2)
add_body(doc,
    "Pretrained backbone parameters (DenseNet-121, ClinicalBERT encoder layers) are "
    "fine-tuned with a learning rate 10× lower than the new components (cross-modal "
    "attention, consistency MLP, classification heads). This prevents catastrophic "
    "forgetting of the pretrained representations while allowing the new layers to "
    "learn quickly."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 5. EVALUATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Evaluation", level=1)

add_heading(doc, "5.1 Metric: AUC-ROC", level=2)
add_body(doc,
    "Following the NIH ChestXray14 benchmark convention, the primary evaluation metric is "
    "the Area Under the ROC Curve (AUC) computed independently for each of the 14 disease "
    "classes. The mean AUC across all classes is reported as the summary metric. AUC is "
    "preferred over accuracy because it is threshold-independent and handles class imbalance "
    "robustly."
)
add_spacer(doc)

add_heading(doc, "5.2 Expected Behaviour", level=2)
add_body(doc,
    "The multimodal model is expected to outperform an image-only DenseNet-121 baseline, "
    "particularly for diseases with strong demographic priors:"
)
add_table(doc,
    ["Disease", "Key Prior Factor", "Expected Gain"],
    [
        ["Cardiomegaly",        "Age, gender",          "High"],
        ["Hernia",              "Age",                  "High"],
        ["Emphysema",           "Age, follow-up count", "Medium"],
        ["Effusion",            "Follow-up visits",     "Medium"],
        ["Pneumothorax",        "View position (AP)",   "Medium"],
        ["Pneumonia",           "Age",                  "Medium"],
        ["Atelectasis",         "View position",        "Low–Medium"],
    ]
)
add_body(doc,
    "The consistency score s provides an additional signal: a low s on a given sample "
    "indicates that the image and history are conflicting (e.g., history says initial "
    "visit but image shows chronic changes), which may correlate with labelling uncertainty."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 6. CODE STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Code Structure", level=1)

add_table(doc,
    ["File / Folder", "Purpose"],
    [
        ["models/model.py",          "Full end-to-end MultimodalDiagnosticModel (forward pass, gated combination)"],
        ["models/visual_encoder.py", "DenseNet-121 and ViT-Base/16 encoders; build_visual_encoder() factory"],
        ["models/text_encoder.py",   "ClinicalHistoryEncoder wrapping Bio_ClinicalBERT; tokenize() helper"],
        ["models/cross_attention.py","CrossModalFusion: bidirectional V→T and T→V attention with projections"],
        ["models/consistency.py",    "ConsistencyModule MLP; consistency_loss(); gated_logits()"],
        ["models/lung_roi.py",       "extract_lung_roi() using torchxrayvision PSPNet segmentation"],
        ["data/dataset.py",          "ChestXray14Dataset; build_clinical_history(); get_transforms()"],
        ["data/preprocess.py",       "NIH_CLASSES list; normalize_labels(); validate_image()"],
        ["train.py",                 "Full training script with differential LR, resume, ROI collate, AUC logging"],
        ["train_simple.py",          "Simplified training loop for readability and quick experimentation"],
        ["infer.py",                 "Single-sample inference with attention heatmap and top-token output"],
        ["demo.py",                  "Presentation-friendly demo: takes patient fields, prints clean results table"],
        ["evaluate.py",              "Batch evaluation on the test split; per-class AUC table"],
        ["app/backend.py",           "Flask/FastAPI backend for the web demo interface"],
        ["app/static/index.html",    "Web demo frontend"],
        ["presentation.html",        "10-minute slide deck (browser-based, arrow-key navigation)"],
    ]
)
add_spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 7. HOW TO RUN
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. How to Run", level=1)

add_heading(doc, "7.1 Training", level=2)
add_code(doc, "python train_simple.py --data_root /path/to/NIH_ChestXray14 --backbone densenet")
add_body(doc, "Full training with all options:")
add_code(doc,
    "python train.py --data_root /path/to/NIH_ChestXray14 \\\n"
    "                --visual_backbone densenet \\\n"
    "                --epochs 50 --batch_size 32 \\\n"
    "                --checkpoint_dir checkpoints/"
)
add_spacer(doc)

add_heading(doc, "7.2 Demo (single image)", level=2)
add_code(doc,
    "python demo.py --image path/to/xray.png \\\n"
    "               --checkpoint checkpoints/best.pth \\\n"
    "               --age 67 --gender M --view PA --followup 2"
)
add_spacer(doc)

add_heading(doc, "7.3 Full Inference with Attention Maps", level=2)
add_code(doc,
    "python infer.py --image path/to/xray.png \\\n"
    "                --history \"Patient is a 67-year-old male. PA view. Follow-up 2.\" \\\n"
    "                --checkpoint checkpoints/best.pth \\\n"
    "                --output_dir results/"
)
add_spacer(doc)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 8. KEY CONTRIBUTIONS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Key Contributions and Novelty", level=1)

add_heading(doc, "8.1 History Conditioning as a First-Class Input", level=2)
add_body(doc,
    "Unlike prior chest X-ray classification models that treat diagnosis as a pure computer "
    "vision problem, this architecture treats patient history as an equal input modality. "
    "The history is not used as a post-hoc filter but participates in the prediction through "
    "learned cross-modal attention weights."
)

add_heading(doc, "8.2 Bidirectional Token-Level Attention", level=2)
add_body(doc,
    "Most multimodal fusion approaches use simple concatenation or late fusion (combining "
    "final-layer embeddings). This model performs attention at the token level in both "
    "directions simultaneously — the image looks at the text and the text looks at the image. "
    "This allows fine-grained spatial-semantic alignment: the cardiac region of the X-ray "
    "can attend to age/gender tokens, and the word 'emphysema' in the history can focus on "
    "the upper lung region."
)

add_heading(doc, "8.3 Consistency Gating", level=2)
add_body(doc,
    "The compatibility score s explicitly models the degree of agreement between the two "
    "modalities. This is important in clinical settings where histories may be incomplete, "
    "incorrect, or contradictory to image findings. A low s causes the model to rely more "
    "heavily on the image-only branch, providing a natural fallback when the history is "
    "unreliable."
)

add_heading(doc, "8.4 Synthetic History for Unimodal Datasets", level=2)
add_body(doc,
    "The NIH ChestXray14 dataset has no free-text clinical notes. Rather than requiring a "
    "new dataset, we demonstrate that structured metadata (age, gender, view, follow-up "
    "number, prior labels) can be converted into natural-language history strings sufficient "
    "to train and benefit from the multimodal architecture. This makes the approach "
    "practically applicable to many existing medical imaging datasets."
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 9. FUTURE WORK
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Future Work", level=1)

add_bullet(doc, "Replace synthetic history strings with real radiology reports from MIMIC-CXR, which pairs 227,000 chest X-rays with actual free-text reports.")
add_bullet(doc, "Extend to additional imaging modalities: CT scans, MRI, and integration with lab results or vital signs.")
add_bullet(doc, "Explainability improvements: visualise both T→V and V→T attention maps side by side; identify which history tokens drive which spatial predictions.")
add_bullet(doc, "Temporal modelling: the current model treats each visit independently. A sequential model over a patient's visit history could better capture disease progression.")
add_bullet(doc, "Clinical validation: prospective evaluation against radiologist readings on a de-identified institutional dataset.")

add_spacer(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 10. REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10. References", level=1)

refs = [
    "Wang, X. et al. (2017). ChestX-ray8: Hospital-scale Chest X-ray Database and Benchmarks. CVPR.",
    "Huang, G. et al. (2017). Densely Connected Convolutional Networks (DenseNet). CVPR.",
    "Rajpurkar, P. et al. (2017). CheXNet: Radiologist-Level Pneumonia Detection on Chest X-Rays with Deep Learning. arXiv:1711.05225.",
    "Alsentzer, E. et al. (2019). Publicly Available Clinical BERT Embeddings. Clinical NLP Workshop @ NAACL.",
    "Devlin, J. et al. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL.",
    "Cohen, J.P. et al. (2022). TorchXRayVision: A library of chest X-ray datasets and models. MIDL.",
    "Vaswani, A. et al. (2017). Attention Is All You Need. NeurIPS.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(ref)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "MV_Project_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
